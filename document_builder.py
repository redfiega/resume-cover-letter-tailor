from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ─────────────────────────────────────────
# STYLE DEFINITIONS
# ─────────────────────────────────────────

STYLES = {
    "Classic": {
        "name_size": 18,
        "header_size": 11,
        "body_size": 12,
        "name_color": RGBColor(0, 0, 0),
        "header_color": RGBColor(0, 0, 0),
        "header_bold": True,
        "header_underline": False,
        "background": False,
        "background_hex": None,
        "margin": 1.0,
        "font": "Times New Roman"
    },
    "Modern": {
        "name_size": 20,
        "header_size": 11,
        "body_size": 12,
        "name_color": RGBColor(0, 101, 164),
        "header_color": RGBColor(0, 101, 164),
        "header_bold": True,
        "header_underline": False,
        "background": False,
        "background_hex": None,
        "margin": 0.75,
        "font": "Calibri"
    },
    "Bold": {
        "name_size": 22,
        "header_size": 12,
        "body_size": 12,
        "name_color": RGBColor(255, 255, 255),
        "header_color": RGBColor(255, 255, 255),
        "header_bold": True,
        "header_underline": False,
        "background": True,
        "background_hex": "1F497D",
        "margin": 0.75,
        "font": "Calibri"
    },
    "Academic": {
        "name_size": 16,
        "header_size": 11,
        "body_size": 12,
        "name_color": RGBColor(0, 0, 0),
        "header_color": RGBColor(0, 0, 0),
        "header_bold": True,
        "header_underline": True,
        "background": False,
        "background_hex": None,
        "margin": 1.0,
        "font": "Georgia"
    }
}


# ─────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────

def set_document_margins(doc: Document, margin_inches: float = 1.0):
    """Set margins for all sections of the document."""
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)


def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph."""
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def add_shading(paragraph, hex_color: str):
    """Add background color shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    clean_color = hex_color.replace("#", "")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), clean_color)
    pPr.append(shd)


def add_inline_text(paragraph, text: str, bold: bool, font_size: int,
                    font_name: str, color: RGBColor = None):
    """Add a text run to a paragraph with specified formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def parse_inline_markdown(paragraph, line: str, font_size: int,
                          font_name: str, color: RGBColor = None):
    """Parse a line with inline **bold** and *italic* markdown."""
    # First handle bold (**text**)
    bold_parts = line.split("**")
    is_bold = False
    for bold_part in bold_parts:
        if bold_part:
            # Now handle italic (*text*) within each bold segment
            italic_parts = bold_part.split("*")
            is_italic = False
            for italic_part in italic_parts:
                if italic_part:
                    run = paragraph.add_run(italic_part)
                    run.bold = is_bold
                    run.italic = is_italic
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                    if color:
                        run.font.color.rgb = color
                is_italic = not is_italic
        is_bold = not is_bold


# ─────────────────────────────────────────
# RESUME BUILDER
# ─────────────────────────────────────────

def build_resume_document(content: str, output_path: str,
                          style_name: str = "Modern") -> str:
    """Convert resume content text into a formatted Word document."""
    style = STYLES.get(style_name, STYLES["Modern"])

    doc = Document()
    set_document_margins(doc, margin_inches=style["margin"])

    lines = content.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):
            # Name / main title
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if style["background"] and style["background_hex"]:
                add_shading(paragraph, style["background_hex"])
            run = paragraph.add_run(line.replace("# ", ""))
            run.bold = True
            run.font.size = Pt(style["name_size"])
            run.font.color.rgb = style["name_color"]
            run.font.name = style["font"]

        elif line.startswith("## "):
            # Section header
            paragraph = doc.add_paragraph()
            if style["background"] and style["background_hex"]:
                add_shading(paragraph, style["background_hex"])
            run = paragraph.add_run(line.replace("## ", "").upper())
            run.bold = style["header_bold"]
            run.underline = style["header_underline"]
            run.font.size = Pt(style["header_size"])
            run.font.color.rgb = style["header_color"]
            run.font.name = style["font"]
            if not style["background"]:
                add_horizontal_line(paragraph)

        elif line.startswith("- ") or line.startswith("* "):
            # Bullet point — handle inline bold
            paragraph = doc.add_paragraph(style="List Bullet")
            parse_inline_markdown(
                paragraph, line[2:],
                style["body_size"], style["font"]
            )

        elif line.startswith("**") and line.endswith("**") and \
                line.count("**") == 2:
            # Fully bold line (job titles, company names)
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.replace("**", ""))
            run.bold = True
            run.font.size = Pt(style["body_size"])
            run.font.name = style["font"]

        else:
            # Regular body text — handle inline bold
            paragraph = doc.add_paragraph()
            parse_inline_markdown(
                paragraph, line,
                style["body_size"], style["font"]
            )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ─────────────────────────────────────────
# COVER LETTER BUILDER
# ─────────────────────────────────────────

def build_cover_letter_document(content: str, output_path: str,
                                style_name: str = "Modern") -> str:
    """Convert cover letter content text into a formatted Word document."""
    style = STYLES.get(style_name, STYLES["Modern"])

    doc = Document()
    set_document_margins(doc, margin_inches=1.0)

    paragraphs = content.split("\n\n")

    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        # Remove markdown headers
        para_text = para_text.replace("## ", "")
        para_text = para_text.replace("# ", "")

        paragraph = doc.add_paragraph()
        # Handle inline bold in cover letter paragraphs
        parse_inline_markdown(
            paragraph, para_text,
            11, style["font"]
        )
        paragraph.space_after = Pt(12)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path